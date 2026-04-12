# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

md_path = r'D:\code\Supermarket Management System\suppermarket system\ss\docs\论文_超市管理系统设计与实现_扩展版.md'
output_path = r'D:\code\Supermarket Management System\suppermarket system\ss\docs\论文_超市管理系统设计与实现_格式版.docx'

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_font(run, font_name='宋体', font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    except:
        pass

def add_title(doc, text, font_size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, '黑体', font_size, True)
    return p

def add_para(doc, text, font_size=12, bold=False, indent=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(2)
    p.alignment = alignment
    run = p.add_run(text)
    set_font(run, '宋体', font_size, bold)
    return p

lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    if line.startswith('# '):
        add_title(doc, line[2:], 22)
    elif line.startswith('## '):
        add_title(doc, line[3:], 16)
    elif line.startswith('### '):
        add_title(doc, line[4:], 14)
    elif line.startswith('#### '):
        add_para(doc, line[5:], 12, True)
    elif line.startswith('**') and line.endswith('**'):
        add_para(doc, line.replace('**', ''), 12, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    elif line.startswith('|') and '|' in line[1:]:
        table_data = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            table_data.append(row)
            i += 1
        if table_data:
            if table_data[0] and all(cell.startswith('---') for cell in table_data[0]):
                table_data = table_data[1:]
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table_data[0]):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    set_font(run, '宋体', 10)
                i -= 1
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        for code_line in code_lines:
            p = doc.add_paragraph(code_line)
            for run in p.runs:
                set_font(run, 'Courier New', 10)
        i -= 1
    elif line.startswith('[') and ']' in line:
        add_para(doc, line, 12)
    elif line.startswith('图'):
        add_para(doc, line, 12)
    elif line.startswith('表'):
        add_para(doc, line, 12)
    else:
        if '**' in line:
            line = line.replace('**', '')
        add_para(doc, line, 12, indent=True)
    
    i += 1

doc.save(output_path)
print(f'已生成: {output_path}')
