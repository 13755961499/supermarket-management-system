# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

# 设置中文字体
def set_chinese_font(run, font_name='宋体', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 读取markdown文件
md_path = r'D:\code\Supermarket Management System\suppermarket system\ss\docs\论文_超市管理系统设计与实现.md'
output_path = r'D:\code\Supermarket Management System\suppermarket system\ss\docs\论文_超市管理系统设计与实现.docx'

with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 解析markdown并添加内容
lines = content.split('\n')
current_paragraph = None

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_chinese_font(run, '黑体', 16 if level == 1 else 14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_paragraph(doc, text, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_chinese_font(run, '宋体', font_size)
        run.bold = bold
    p.alignment = alignment
    return p

# 处理每一行
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    if not line:
        i += 1
        continue
    
    # 标题处理
    if line.startswith('# '):
        # 论文标题
        add_paragraph(doc, line[2:], font_size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    elif line.startswith('## '):
        # 章标题
        add_heading(doc, line[3:], level=1)
    elif line.startswith('### '):
        # 节标题
        add_heading(doc, line[4:], level=2)
    elif line.startswith('#### '):
        # 小节标题
        add_heading(doc, line[5:], level=3)
    elif line.startswith('**') and line.endswith('**'):
        # 关键词
        add_paragraph(doc, line.replace('**', ''), font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    elif line.startswith('|') and '|' in line[1:]:
        # 表格
        table_data = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            table_data.append(row)
            i += 1
        if table_data:
            # 确定列数
            max_cols = max(len(row) for row in table_data)
            # 跳过分隔符行
            if table_data[0] and all(cell.startswith('---') for cell in table_data[0]):
                table_data = table_data[1:]
            
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table_data[0]):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    set_chinese_font(run, '宋体', 10)
                i -= 1  # 回退一行，因为外层循环会i+=1
    elif line.startswith('```'):
        # 代码块
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        for code_line in code_lines:
            p = doc.add_paragraph(code_line)
            for run in p.runs:
                set_chinese_font(run, 'Courier New', 10)
        i -= 1
    elif line.startswith('[1]') or line.startswith('['):
        # 参考文献
        add_paragraph(doc, line, font_size=12)
    elif line.startswith('*本论文约'):
        # 字数统计
        pass
    elif '附录' in line:
        # 附录
        add_heading(doc, line, level=1)
    else:
        # 正文
        # 处理加粗文本
        text = line
        if '**' in text:
            # 简单处理：如果有**，添加普通段落
            text = text.replace('**', '')
        add_paragraph(doc, text, font_size=12)
    
    i += 1

# 保存文档
doc.save(output_path)
print(f'论文已生成: {output_path}')
